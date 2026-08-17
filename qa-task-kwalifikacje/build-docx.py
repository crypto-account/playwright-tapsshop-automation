"""
Konwersja zadanie-qa-kwalifikacje.md -> zadanie-qa-kwalifikacje.docx
Format zoptymalizowany pod import do Google Docs (zachowuje tabele, nagłówki, pogrubienie).
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
SRC = HERE / "zadanie-qa-kwalifikacje.md"
DST = HERE / "zadanie-qa-kwalifikacje.docx"


def add_bold_runs(paragraph, text):
    """Zamienia **bold** na wytłuszczone runy w akapicie. Obsługuje też `code`."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def parse_md_table(lines, start_idx):
    """Zwraca (rows, next_idx). rows[0] = nagłówek."""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        # pomiń separator |---|---|
        if re.match(r"^\|[\s\-:|]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            # tekst może zawierać <br> - zamień na akapity
            segments = cell_text.split("<br>")
            for s_idx, seg in enumerate(segments):
                if s_idx == 0:
                    para = cell.paragraphs[0]
                else:
                    para = cell.add_paragraph()
                add_bold_runs(para, seg.strip())
                # nagłówek na bold
                if r_idx == 0:
                    for run in para.runs:
                        run.bold = True
            if r_idx == 0:
                set_cell_shading(cell, "D9E2F3")


def build():
    doc = Document()

    # domyślny styl
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ustaw język PL na dokumencie
    styles_element = doc.styles.element
    rpr_default = styles_element.xpath("./w:docDefaults/w:rPrDefault/w:rPr")
    if rpr_default:
        lang = OxmlElement("w:lang")
        lang.set(qn("w:val"), "pl-PL")
        rpr_default[0].append(lang)

    # marginesy
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    md = SRC.read_text(encoding="utf-8")
    lines = md.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # tabela
        if stripped.startswith("|") and "|" in stripped[1:]:
            rows, i = parse_md_table(lines, i)
            add_table(doc, rows)
            doc.add_paragraph()
            continue

        # separator ---
        if stripped == "---":
            i += 1
            continue

        # obraz markdown ![alt](path)
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if img_match:
            alt_text, img_path = img_match.groups()
            img_full = HERE / img_path
            if img_full.exists():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                # szerokość obrazu 16 cm (mieści się w marginesach A4 z 2 cm)
                run.add_picture(str(img_full), width=Cm(16))
                # podpis pod obrazem (mniejszy, italic)
                if alt_text:
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption.add_run(alt_text)
                    caption_run.italic = True
                    caption_run.font.size = Pt(9)
            else:
                para = doc.add_paragraph()
                para.add_run(f"[BRAK OBRAZKA: {img_path}]").italic = True
            i += 1
            continue

        # nagłówki
        if stripped.startswith("# "):
            h = doc.add_heading(stripped[2:].strip(), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        # numerowana lista "1. ..." / "2. ..."
        elif re.match(r"^\d+\.\s", stripped):
            para = doc.add_paragraph(style="List Number")
            text = re.sub(r"^\d+\.\s", "", stripped)
            add_bold_runs(para, text)
        # punktowana lista
        elif stripped.startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            add_bold_runs(para, stripped[2:])
        # kontynuacja listy (wcięcie)
        elif line.startswith("   ") and stripped:
            para = doc.add_paragraph(style="List Bullet 2")
            add_bold_runs(para, stripped.lstrip("- "))
        elif stripped == "":
            # pusta linia — nic nie dodawaj (odstępy robią nagłówki/tabele)
            pass
        else:
            para = doc.add_paragraph()
            add_bold_runs(para, stripped)

        i += 1

    doc.save(DST)
    print(f"OK  {DST.name}  ({DST.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    build()
